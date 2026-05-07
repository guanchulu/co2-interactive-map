from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SUBMISSION = ROOT / "docs" / "joule_submission"
MANUSCRIPT = SUBMISSION / "manuscript.md"
CAPTIONS = SUBMISSION / "figure_captions.md"
FIGURE_DIR = SUBMISSION / "figures_storyline"
WORD_FIGURES = SUBMISSION / "figures_word"
OUT = SUBMISSION / "co2_joule_manuscript.docx"


IMAGE_RE = re.compile(r"!\[(?P<label>[^\]]*)\]\((?P<path>[^)]+)\)")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")


def set_document_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    for style_name, size in {
        "Title": 16,
        "Heading 1": 14,
        "Heading 2": 13,
        "Heading 3": 12,
    }.items():
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_inline_markdown(paragraph, text: str) -> None:
    """Add a small Markdown subset: bold, inline code, and links."""
    text = LINK_RE.sub(r"\1 (\2)", text)
    pos = 0
    for match in re.finditer(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_captions(path: Path) -> dict[str, str]:
    captions: dict[str, str] = {}
    current_key: str | None = None
    current_title: str | None = None
    current_lines: list[str] = []

    def flush() -> None:
        nonlocal current_key, current_title, current_lines
        if current_key and current_title:
            body = " ".join(line.strip() for line in current_lines if line.strip())
            captions[current_key] = f"{current_title}. {body}".strip()
        current_key = None
        current_title = None
        current_lines = []

    for raw in path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        match = re.match(r"##\s+(Figure\s+(\d+)\.\s+.+)", line)
        if match:
            flush()
            current_key = f"Figure {match.group(2)}"
            current_title = match.group(1)
        elif line.startswith("## "):
            flush()
        elif current_key:
            current_lines.append(line)
    flush()
    return captions


def convert_svg_to_png(svg: Path, png: Path) -> Path | None:
    if not svg.exists():
        return None
    try:
        import cairosvg

        png.parent.mkdir(parents=True, exist_ok=True)
        cairosvg.svg2png(url=str(svg), write_to=str(png), output_width=2200)
        return png
    except Exception as exc:  # pragma: no cover - graceful export fallback
        print(f"Could not convert {svg.name}: {exc}")
        return None


def figure_png_for(markdown_path: str) -> Path | None:
    figure_path = (SUBMISSION / markdown_path).resolve()
    if figure_path.suffix.lower() == ".svg":
        png = WORD_FIGURES / f"{figure_path.stem}.png"
        return convert_svg_to_png(figure_path, png)
    if figure_path.exists():
        return figure_path
    return None


def add_figure(doc: Document, label: str, figure_path: str, captions: dict[str, str]) -> None:
    figure_num_match = re.search(r"Figure\s+(\d+)", label)
    figure_key = f"Figure {figure_num_match.group(1)}" if figure_num_match else label
    png = figure_png_for(figure_path)

    if png and png.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(png), width=Inches(6.4))
    else:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(f"[{label}: {figure_path}]")
        run.italic = True

    caption_text = captions.get(figure_key)
    if caption_text:
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(caption_text)
        run.italic = True


def flush_paragraph(doc: Document, buffer: list[str]) -> None:
    if not buffer:
        return
    text = " ".join(part.strip() for part in buffer if part.strip())
    if text:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(6)
        add_inline_markdown(paragraph, text)
    buffer.clear()


def markdown_to_docx(markdown: Path, out: Path) -> None:
    doc = Document()
    set_document_styles(doc)
    captions = parse_captions(CAPTIONS)
    paragraph_buffer: list[str] = []

    lines = markdown.read_text(encoding="utf-8").splitlines()
    for raw_line in lines:
        line = raw_line.rstrip()

        if not line.strip():
            flush_paragraph(doc, paragraph_buffer)
            continue

        image_match = IMAGE_RE.match(line.strip())
        if image_match:
            flush_paragraph(doc, paragraph_buffer)
            add_figure(
                doc,
                image_match.group("label"),
                image_match.group("path"),
                captions,
            )
            continue

        if line.startswith("# "):
            flush_paragraph(doc, paragraph_buffer)
            title = line[2:].strip()
            paragraph = doc.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_markdown(paragraph, title)
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading_match:
            flush_paragraph(doc, paragraph_buffer)
            level = min(len(heading_match.group(1)) - 1, 3)
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            add_inline_markdown(paragraph, heading_match.group(2).strip())
            continue

        bullet_match = re.match(r"^[-*]\s+(.+)$", line)
        if bullet_match:
            flush_paragraph(doc, paragraph_buffer)
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, bullet_match.group(1).strip())
            continue

        paragraph_buffer.append(line)

    flush_paragraph(doc, paragraph_buffer)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)


if __name__ == "__main__":
    markdown_to_docx(MANUSCRIPT, OUT)
    print(OUT)
